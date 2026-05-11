"""
MCP Docx Processing Service
Provides a single gateway tool "docx_process" that routes to document operations.
Implemented using the official MCP library
"""

import os
import json
import tempfile
import time
import logging
from contextlib import asynccontextmanager
from typing import AsyncIterator, Dict, Any

from mcp.server.fastmcp import FastMCP, Context
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING, WD_COLOR_INDEX
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    handlers=[
        logging.FileHandler(os.path.join(tempfile.gettempdir(), "docx_mcp_server.log")),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger("DocxMCPServer")

class DocxProcessor:
    """Class for processing Docx documents, implementing various document operations"""

    def __init__(self):
        self.documents = {}  # Store opened documents
        self.current_document = None
        self.current_file_path = None

# Create global processor instance
processor = DocxProcessor()

@asynccontextmanager
async def server_lifespan(server: FastMCP) -> AsyncIterator[Dict[str, Any]]:
    """Manage server lifecycle"""
    try:
        logger.info("DocxProcessor MCP server starting with clean state...")
        yield {"processor": processor}
    finally:
        logger.info("DocxProcessor MCP server shutting down...")

# Create MCP server
mcp = FastMCP(
    name="DocxProcessor",
    instructions="Word document processing service. Use the docx_process gateway tool with route and params to perform document operations. Refer to the docx-process skill for available routes and parameter details.",
    lifespan=server_lifespan
)

# ============================================================================
# Helper functions
# ============================================================================

def _save_paragraph_styles(doc, start_index, count):
    """Save style and format info for paragraphs in range [start_index, start_index+count)"""
    original_styles = []
    for i in range(start_index, start_index + count):
        if i < len(doc.paragraphs):
            para = doc.paragraphs[i]
            style_info = {
                'style': para.style,
                'alignment': para.alignment,
                'runs': []
            }
            for run in para.runs:
                run_info = {
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    'font_size': run.font.size,
                    'font_name': run.font.name,
                    'color': run.font.color.rgb if run.font.color.rgb else None
                }
                style_info['runs'].append(run_info)
            original_styles.append(style_info)
        else:
            original_styles.append(original_styles[-1] if original_styles else {
                'style': None, 'alignment': None, 'runs': []
            })
    while len(original_styles) < count:
        original_styles.append(original_styles[-1] if original_styles else {
            'style': None, 'alignment': None, 'runs': []
        })
    return original_styles


def _replace_paragraphs_in_range(doc, start_index, end_index, new_content, original_styles):
    """Delete paragraphs in [start_index, end_index) and insert new_content with original styles"""
    insert_position = start_index

    for i in range(end_index - 1, start_index - 1, -1):
        p = doc.paragraphs[i]._element
        p.getparent().remove(p)

    for i, content in enumerate(reversed(new_content)):
        p = doc.add_paragraph()
        style_info = original_styles[len(new_content) - i - 1]
        if style_info['style']:
            p.style = style_info['style']
        if style_info['alignment'] is not None:
            p.alignment = style_info['alignment']

        if style_info['runs']:
            run = p.add_run(content)
            run_info = style_info['runs'][0]
            run.bold = run_info['bold']
            run.italic = run_info['italic']
            run.underline = run_info['underline']
            if run_info['font_size']:
                run.font.size = run_info['font_size']
            if run_info['font_name']:
                _set_run_font(run, run_info['font_name'])
            if run_info['color']:
                run.font.color.rgb = run_info['color']
        else:
            p.text = content

        doc._body._body.insert(insert_position, p._p)
        doc._body._body.remove(doc.paragraphs[-1]._p)


def _get_lock_file_path(file_path: str) -> str:
    """Get the Word lock file path for a given document path.
    Word creates ~$filename.docx when it has a file open."""
    directory = os.path.dirname(file_path)
    basename = os.path.basename(file_path)
    return os.path.join(directory, f"~${basename}")


def _is_file_locked_by_word(file_path: str) -> bool:
    """Check if a file appears to be locked by Microsoft Word.
    Detects the hidden ~$ lock file that Word creates when a document is open."""
    lock_file = _get_lock_file_path(file_path)
    return os.path.exists(lock_file)


def _set_run_font(run, font_name: str):
    """Set font name on a run, including eastAsia for CJK support."""
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


def _parse_hex_color(color: str) -> RGBColor:
    """Parse a #RRGGBB hex color string to RGBColor. Raises ValueError on invalid format."""
    if not color.startswith('#') or len(color) != 7:
        raise ValueError(f"Invalid color format: '{color}'. Use '#RRGGBB' format (e.g. '#FF0000')")
    r = int(color[1:3], 16)
    g = int(color[3:5], 16)
    b = int(color[5:7], 16)
    return RGBColor(r, g, b)


# ============================================================================
# Route handler functions
# ============================================================================

def _handle_create_document(params: dict) -> str:
    file_path = params.get("file_path")
    if not file_path:
        return "Missing required parameter: file_path"
    try:
        processor.current_document = Document()
        processor.current_file_path = file_path
        processor.documents[file_path] = processor.current_document
        processor.current_document.save(file_path)
        return f"Document created successfully: {file_path}"
    except Exception as e:
        return f"Failed to create document: {str(e)}"


def _handle_open_document(params: dict) -> str:
    file_path = params.get("file_path")
    if not file_path:
        return "Missing required parameter: file_path"
    try:
        if not os.path.exists(file_path):
            return f"File does not exist: {file_path}"
        if processor.current_document and processor.current_file_path and processor.current_file_path != file_path:
            return (
                f"Another document is already open: {processor.current_file_path}. "
                "Please close it first using close_document before opening a new one."
            )
        if _is_file_locked_by_word(file_path):
            return (
                f"File is currently open in Microsoft Word: {file_path}. "
                "Please close the file in Word before opening it here."
            )
        processor.current_document = Document(file_path)
        processor.current_file_path = file_path
        processor.documents[file_path] = processor.current_document
        return f"Document opened successfully: {file_path}"
    except PermissionError:
        return (
            f"Cannot open file (permission denied): {file_path}. "
            "The file may be open in Microsoft Word. Please close it and try again."
        )
    except Exception as e:
        return f"Failed to open document: {str(e)}"


def _handle_save_document(params: dict) -> str:
    try:
        if not processor.current_document:
            return "No document is open"
        if not processor.current_file_path:
            return "Current document has not been saved before, please use save_as_document to specify a save path"
        save_path = processor.current_file_path
        if _is_file_locked_by_word(save_path):
            return (
                f"Cannot save: the file is open in Microsoft Word: {save_path}. "
                "Please close the file in Word before saving."
            )
        last_error = None
        for attempt in range(3):
            try:
                processor.current_document.save(save_path)
                return f"Document saved successfully to original file: {save_path}"
            except PermissionError as e:
                last_error = e
                if attempt < 2:
                    time.sleep(2)
        return (
            f"Failed to save document after 3 attempts: {save_path}. "
            "The file may be open in Microsoft Word. Please close it and try again. "
            f"Last error: {last_error}"
        )
    except Exception as e:
        return f"Failed to save document: {str(e)}"


def _handle_add_paragraph(params: dict) -> str:
    text = params.get("text")
    if text is None:
        return "Missing required parameter: text"
    bold = params.get("bold", False)
    italic = params.get("italic", False)
    underline = params.get("underline", False)
    font_size = params.get("font_size")
    font_name = params.get("font_name")
    color = params.get("color")
    alignment = params.get("alignment")
    try:
        if not processor.current_document:
            return "No document is open"
        paragraph = processor.current_document.add_paragraph(text)
        if paragraph.runs:
            run = paragraph.runs[0]
            run.bold = bold
            run.italic = italic
            run.underline = underline
            if font_size:
                run.font.size = Pt(font_size)
            if font_name:
                _set_run_font(run, font_name)
            if color:
                if not color.startswith('#') or len(color) != 7:
                    return f"Invalid color format: '{color}'. Use '#RRGGBB' format (e.g. '#FF0000')"
                run.font.color.rgb = _parse_hex_color(color)
        if alignment:
            if alignment == "left":
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            elif alignment == "center":
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif alignment == "right":
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            elif alignment == "justify":
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        return "Paragraph added"
    except Exception as e:
        return f"Failed to add paragraph: {str(e)}"


def _handle_add_heading(params: dict) -> str:
    text = params.get("text")
    level = params.get("level")
    if text is None:
        return "Missing required parameter: text"
    if level is None:
        return "Missing required parameter: level"
    if not isinstance(level, int) or not (1 <= level <= 9):
        return "Invalid level: must be an integer between 1 and 9"
    try:
        if not processor.current_document:
            return "No document is open"
        processor.current_document.add_heading(text, level=level)
        return f"Added level {level} heading"
    except Exception as e:
        return f"Failed to add heading: {str(e)}"


def _handle_add_table(params: dict) -> str:
    rows = params.get("rows")
    cols = params.get("cols")
    data = params.get("data")
    if rows is None or cols is None:
        return "Missing required parameters: rows, cols"
    if not isinstance(rows, int) or rows < 1:
        return "Invalid rows: must be a positive integer"
    if not isinstance(cols, int) or cols < 1:
        return "Invalid cols: must be a positive integer"
    try:
        if not processor.current_document:
            return "No document is open"
        table = processor.current_document.add_table(rows=rows, cols=cols, style="Table Grid")
        if data:
            for i, row_data in enumerate(data):
                if i < rows:
                    row = table.rows[i]
                    for j, cell_text in enumerate(row_data):
                        if j < cols:
                            row.cells[j].text = str(cell_text)
        return f"Added {rows}x{cols} table"
    except Exception as e:
        return f"Failed to add table: {str(e)}"


def _handle_get_document_info(params: dict) -> str:
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        sections_count = len(doc.sections)
        paragraphs_count = len(doc.paragraphs)
        tables_count = len(doc.tables)
        paragraph_styles = []
        for style in doc.styles:
            if style.type == WD_STYLE_TYPE.PARAGRAPH:
                paragraph_styles.append(style.name)
        info = f"Document path: {processor.current_file_path}\n"
        info += f"Section count: {sections_count}\n"
        info += f"Paragraph count: {paragraphs_count}\n"
        info += f"Table count: {tables_count}\n"
        info += f"Available paragraph styles: {', '.join(paragraph_styles[:10])}..."
        return info
    except Exception as e:
        return f"Failed to get document information: {str(e)}"


def _handle_search_text(params: dict) -> str:
    keyword = params.get("keyword")
    if keyword is None:
        return "Missing required parameter: keyword"
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        results = []
        for i, paragraph in enumerate(doc.paragraphs):
            if keyword in paragraph.text:
                results.append({"type": "paragraph", "index": i, "text": paragraph.text})
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    if keyword in cell.text:
                        results.append({
                            "type": "table cell",
                            "table_index": t_idx,
                            "row": r_idx,
                            "column": c_idx,
                            "text": cell.text
                        })
        if not results:
            return f"Keyword '{keyword}' not found"
        response = f"Found {len(results)} occurrences of '{keyword}':\n\n"
        for idx, result in enumerate(results):
            response += f"{idx+1}. {result['type']} "
            if result['type'] == "paragraph":
                response += f"index {result['index']}: {result['text'][:100]}"
                if len(result['text']) > 100:
                    response += "..."
                response += "\n"
            else:
                response += f"in table {result['table_index']} at cell ({result['row']},{result['column']}): {result['text'][:100]}"
                if len(result['text']) > 100:
                    response += "..."
                response += "\n"
        return response
    except Exception as e:
        return f"Failed to search text: {str(e)}"


def _handle_search_and_replace(params: dict) -> str:
    keyword = params.get("keyword")
    replace_with = params.get("replace_with")
    preview_only = params.get("preview_only", False)
    if keyword is None or replace_with is None:
        return "Missing required parameters: keyword, replace_with"
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        results = []
        for i, paragraph in enumerate(doc.paragraphs):
            if keyword in paragraph.text:
                original_text = paragraph.text
                replaced_text = original_text.replace(keyword, replace_with)
                results.append({
                    "type": "paragraph", "index": i,
                    "original": original_text, "replaced": replaced_text,
                    "count": original_text.count(keyword)
                })
                if not preview_only:
                    paragraph.text = replaced_text
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    if keyword in cell.text:
                        original_text = cell.text
                        replaced_text = original_text.replace(keyword, replace_with)
                        results.append({
                            "type": "table cell", "table_index": t_idx,
                            "row": r_idx, "column": c_idx,
                            "original": original_text, "replaced": replaced_text,
                            "count": original_text.count(keyword)
                        })
                        if not preview_only:
                            for para in cell.paragraphs:
                                if keyword in para.text:
                                    para.text = para.text.replace(keyword, replace_with)
        if not results:
            return f"Keyword '{keyword}' not found"
        total_replacements = sum(item["count"] for item in results)
        action_word = "Preview" if preview_only else "Replace"
        response = f"{action_word} '{keyword}' with '{replace_with}', found {len(results)} locations, {total_replacements} occurrences:\n\n"
        for idx, result in enumerate(results):
            response += f"{idx+1}. In {result['type']} "
            if result['type'] == "paragraph":
                response += f"index {result['index']} {action_word.lower()}ing {result['count']} times:\n"
            else:
                response += f"table {result['table_index']} at cell ({result['row']},{result['column']}) {action_word.lower()}ing {result['count']} times:\n"
            max_display = 50
            if len(result['original']) > max_display * 2:
                start_pos = result['original'].find(keyword)
                start_pos = max(0, start_pos - max_display)
                excerpt_original = "..." + result['original'][start_pos:start_pos + max_display * 2] + "..."
                excerpt_replaced = "..." + result['replaced'][start_pos:start_pos + max_display * 2] + "..."
            else:
                excerpt_original = result['original']
                excerpt_replaced = result['replaced']
            response += f"  Original: {excerpt_original}\n"
            response += f"  Replaced: {excerpt_replaced}\n\n"
        if preview_only:
            response += "This is a preview of replacements. No actual changes were made. To execute replacements, set preview_only to False."
        else:
            response += "Replacements completed successfully."
        return response
    except Exception as e:
        return f"Search and replace failed: {str(e)}"


def _handle_find_and_replace(params: dict) -> str:
    find_text = params.get("find_text")
    replace_text = params.get("replace_text")
    if find_text is None or replace_text is None:
        return "Missing required parameters: find_text, replace_text"
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        replace_count = 0
        for paragraph in doc.paragraphs:
            if find_text in paragraph.text:
                count = paragraph.text.count(find_text)
                paragraph.text = paragraph.text.replace(find_text, replace_text)
                replace_count += count
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if find_text in paragraph.text:
                            count = paragraph.text.count(find_text)
                            paragraph.text = paragraph.text.replace(find_text, replace_text)
                            replace_count += count
        return f"Replaced '{find_text}' with '{replace_text}', {replace_count} occurrences"
    except Exception as e:
        return f"Find and replace failed: {str(e)}"


def _handle_merge_table_cells(params: dict) -> str:
    table_index = params.get("table_index")
    start_row = params.get("start_row")
    start_col = params.get("start_col")
    end_row = params.get("end_row")
    end_col = params.get("end_col")
    if table_index is None or start_row is None or start_col is None or end_row is None or end_col is None:
        return "Missing required parameters: table_index, start_row, start_col, end_row, end_col"
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        if not doc.tables:
            return "No tables in document"
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Table index out of range: {table_index}, document has {len(doc.tables)} tables"
        table = doc.tables[table_index]
        if start_row < 0 or start_row >= len(table.rows):
            return f"Start row index out of range: {start_row}, table has {len(table.rows)} rows"
        if start_col < 0 or start_col >= len(table.columns):
            return f"Start column index out of range: {start_col}, table has {len(table.columns)} columns"
        if end_row < start_row or end_row >= len(table.rows):
            return f"End row index invalid: {end_row}, should be between {start_row} and {len(table.rows)-1}"
        if end_col < start_col or end_col >= len(table.columns):
            return f"End column index invalid: {end_col}, should be between {start_col} and {len(table.columns)-1}"
        start_cell = table.cell(start_row, start_col)
        end_cell = table.cell(end_row, end_col)
        start_cell.merge(end_cell)
        return f"Merged cells in table {table_index} from ({start_row},{start_col}) to ({end_row},{end_col})"
    except Exception as e:
        return f"Failed to merge table cells: {str(e)}"


def _handle_split_table(params: dict) -> str:
    table_index = params.get("table_index")
    row_index = params.get("row_index")
    if table_index is None or row_index is None:
        return "Missing required parameters: table_index, row_index"
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        if not doc.tables:
            return "No tables in document"
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Table index out of range: {table_index}, document has {len(doc.tables)} tables"
        table = doc.tables[table_index]
        if row_index < 0 or row_index >= len(table.rows) - 1:
            return f"Row index invalid: {row_index}, should be between 0 and {len(table.rows)-2}"
        tbl = table._tbl
        split_position = row_index + 1
        new_tbl = OxmlElement('w:tbl')
        for child in tbl.xpath('./w:tblPr')[0].getchildren():
            new_tbl.append(child.copy())
        for child in tbl.xpath('./w:tblGrid')[0].getchildren():
            new_tbl.append(child.copy())
        rows = tbl.xpath('./w:tr')
        for i in range(split_position, len(rows)):
            new_tbl.append(rows[i])
        tbl.addnext(new_tbl)
        return f"Split table {table_index} after row {row_index}"
    except Exception as e:
        return f"Failed to split table: {str(e)}"


def _handle_add_table_row(params: dict) -> str:
    table_index = params.get("table_index")
    data = params.get("data")
    if table_index is None:
        return "Missing required parameter: table_index"
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        if not doc.tables:
            return "No tables in document"
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Table index out of range: {table_index}, document has {len(doc.tables)} tables"
        table = doc.tables[table_index]
        new_row = table.add_row()
        if data:
            for i, cell_text in enumerate(data):
                if i < len(new_row.cells):
                    new_row.cells[i].text = str(cell_text)
        return f"Added new row to table {table_index}"
    except Exception as e:
        return f"Failed to add table row: {str(e)}"


def _handle_delete_table_row(params: dict) -> str:
    table_index = params.get("table_index")
    row_index = params.get("row_index")
    if table_index is None or row_index is None:
        return "Missing required parameters: table_index, row_index"
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        if not doc.tables:
            return "No tables in document"
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Table index out of range: {table_index}, document has {len(doc.tables)} tables"
        table = doc.tables[table_index]
        if row_index < 0 or row_index >= len(table.rows):
            return f"Row index out of range: {row_index}, table has {len(table.rows)} rows"
        row = table.rows[row_index]._tr
        row.getparent().remove(row)
        return f"Deleted row {row_index} from table {table_index}"
    except Exception as e:
        return f"Failed to delete table row: {str(e)}"


def _handle_edit_table_cell(params: dict) -> str:
    table_index = params.get("table_index")
    row_index = params.get("row_index")
    col_index = params.get("col_index")
    text = params.get("text")
    if table_index is None or row_index is None or col_index is None or text is None:
        return "Missing required parameters: table_index, row_index, col_index, text"
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        if not doc.tables:
            return "No tables in document"
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Table index out of range: {table_index}, document has {len(doc.tables)} tables"
        table = doc.tables[table_index]
        if row_index < 0 or row_index >= len(table.rows):
            return f"Row index out of range: {row_index}, table has {len(table.rows)} rows"
        if col_index < 0 or col_index >= len(table.columns):
            return f"Column index out of range: {col_index}, table has {len(table.columns)} columns"
        table.cell(row_index, col_index).text = text
        return f"Cell ({row_index}, {col_index}) in table {table_index} has been modified"
    except Exception as e:
        return f"Failed to edit table cell: {str(e)}"


def _handle_add_page_break(params: dict) -> str:
    try:
        if not processor.current_document:
            return "No document is open"
        processor.current_document.add_page_break()
        return "Page break added"
    except Exception as e:
        return f"Failed to add page break: {str(e)}"


def _handle_set_page_margins(params: dict) -> str:
    top = params.get("top")
    bottom = params.get("bottom")
    left = params.get("left")
    right = params.get("right")
    section_index = params.get("section_index", 0)
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        if section_index < 0 or section_index >= len(doc.sections):
            return f"Section index out of range: {section_index}, document has {len(doc.sections)} sections"
        section = doc.sections[section_index]
        if top is not None:
            section.top_margin = Cm(top)
        if bottom is not None:
            section.bottom_margin = Cm(bottom)
        if left is not None:
            section.left_margin = Cm(left)
        if right is not None:
            section.right_margin = Cm(right)
        return f"Page margins set for section {section_index}"
    except Exception as e:
        return f"Failed to set page margins: {str(e)}"


def _handle_delete_paragraph(params: dict) -> str:
    paragraph_index = params.get("paragraph_index")
    if paragraph_index is None:
        return "Missing required parameter: paragraph_index"
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Paragraph index out of range: {paragraph_index}, document has {len(doc.paragraphs)} paragraphs"
        paragraph = doc.paragraphs[paragraph_index]
        p = paragraph._element
        p.getparent().remove(p)
        paragraph._p = None
        paragraph._element = None
        return f"Paragraph {paragraph_index} deleted"
    except Exception as e:
        return f"Failed to delete paragraph: {str(e)}"


def _handle_delete_text(params: dict) -> str:
    paragraph_index = params.get("paragraph_index")
    start_pos = params.get("start_pos")
    end_pos = params.get("end_pos")
    if paragraph_index is None or start_pos is None or end_pos is None:
        return "Missing required parameters: paragraph_index, start_pos, end_pos"
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Paragraph index out of range: {paragraph_index}, document has {len(doc.paragraphs)} paragraphs"
        paragraph = doc.paragraphs[paragraph_index]
        text = paragraph.text
        if start_pos < 0 or start_pos >= len(text):
            return f"Start position out of range: {start_pos}, paragraph length is {len(text)}"
        if end_pos <= start_pos or end_pos > len(text):
            return f"End position invalid: {end_pos}, should be between {start_pos+1} and {len(text)}"
        new_text = text[:start_pos] + text[end_pos:]
        paragraph.text = new_text
        return f"Deleted text from position {start_pos} to {end_pos} in paragraph {paragraph_index}"
    except Exception as e:
        return f"Failed to delete text: {str(e)}"


def _handle_save_as_document(params: dict) -> str:
    new_file_path = params.get("new_file_path")
    if not new_file_path:
        return "Missing required parameter: new_file_path"
    try:
        if not processor.current_document:
            return "No document is open"
        if _is_file_locked_by_word(new_file_path):
            return (
                f"Cannot save: the target file is open in Microsoft Word: {new_file_path}. "
                "Please close it in Word before saving."
            )
        processor.current_document.save(new_file_path)
        old_path = processor.current_file_path
        if old_path and old_path in processor.documents:
            del processor.documents[old_path]
        processor.current_file_path = new_file_path
        processor.documents[new_file_path] = processor.current_document
        return f"Document saved as: {new_file_path}"
    except PermissionError:
        return (
            f"Failed to save document (permission denied): {new_file_path}. "
            "The target file may be open in Microsoft Word. Please close it and try again."
        )
    except Exception as e:
        return f"Failed to save document: {str(e)}"


def _handle_create_document_copy(params: dict) -> str:
    suffix = params.get("suffix", "-副本")
    try:
        if not processor.current_document:
            return "No document is open"
        if not processor.current_file_path:
            return "Current document has not been saved, cannot create a copy"
        file_dir = os.path.dirname(processor.current_file_path)
        file_name = os.path.basename(processor.current_file_path)
        file_name_without_ext, file_ext = os.path.splitext(file_name)
        new_file_name = f"{file_name_without_ext}{suffix}{file_ext}"
        new_file_path = os.path.join(file_dir, new_file_name)
        if _is_file_locked_by_word(new_file_path):
            return (
                f"Cannot create copy: the target file is open in Microsoft Word: {new_file_path}. "
                "Please close it in Word before creating a copy."
            )
        processor.current_document.save(new_file_path)
        return f"Document copy created: {new_file_path}"
    except PermissionError:
        return (
            f"Failed to create document copy (permission denied). "
            "The target file may be open in Microsoft Word. Please close it and try again."
        )
    except Exception as e:
        return f"Failed to create document copy: {str(e)}"


def _handle_close_document(params: dict) -> str:
    if processor.current_file_path and processor.current_file_path in processor.documents:
        del processor.documents[processor.current_file_path]
    processor.current_document = None
    processor.current_file_path = None
    return "Document closed"


def _handle_reload_document(params: dict) -> str:
    if not processor.current_file_path:
        return "No document open"
    file_path = processor.current_file_path
    if not os.path.exists(file_path):
        return f"File no longer exists: {file_path}"
    try:
        if _is_file_locked_by_word(file_path):
            return (
                f"Cannot reload: the file is open in Microsoft Word: {file_path}. "
                "Please close the file in Word before reloading."
            )
        processor.current_document = Document(file_path)
        return f"Document reloaded from: {file_path}"
    except PermissionError:
        return (
            f"Cannot reload document (permission denied): {file_path}. "
            "The file may be open in Microsoft Word. Please close it and try again."
        )
    except Exception as e:
        return f"Failed to reload document: {str(e)}"


def _handle_replace_section(params: dict) -> str:
    section_title = params.get("section_title")
    new_content = params.get("new_content")
    preserve_title = params.get("preserve_title", True)
    if section_title is None or new_content is None:
        return "Missing required parameters: section_title, new_content"
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        title_index = -1
        for i, paragraph in enumerate(doc.paragraphs):
            if section_title in paragraph.text:
                title_index = i
                break
        if title_index == -1:
            return f"Title not found: '{section_title}'"
        end_index = len(doc.paragraphs)
        title_style = doc.paragraphs[title_index].style
        def _heading_level(style):
            if style and style.name and style.name.startswith('Heading'):
                try:
                    return int(style.name.split()[-1])
                except ValueError:
                    pass
            return None
        title_level = _heading_level(title_style)
        for i in range(title_index + 1, len(doc.paragraphs)):
            other_level = _heading_level(doc.paragraphs[i].style)
            if other_level is not None and (title_level is None or other_level <= title_level):
                end_index = i
                break
        start_delete = title_index + (1 if preserve_title else 0)
        original_styles = _save_paragraph_styles(doc, start_delete, len(new_content))
        _replace_paragraphs_in_range(doc, start_delete, end_index, new_content, original_styles)
        return f"Replaced content under title '{section_title}', keeping original format and style"
    except Exception as e:
        return f"Failed to replace content: {str(e)}"


def _handle_edit_section_by_keyword(params: dict) -> str:
    keyword = params.get("keyword")
    new_content = params.get("new_content")
    section_range = params.get("section_range", 3)
    if keyword is None or new_content is None:
        return "Missing required parameters: keyword, new_content"
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        keyword_indices = []
        for i, paragraph in enumerate(doc.paragraphs):
            if keyword in paragraph.text:
                keyword_indices.append(i)
        if not keyword_indices:
            return f"Keyword not found: '{keyword}'"
        keyword_index = keyword_indices[0]
        start_index = max(0, keyword_index - section_range)
        end_index = min(len(doc.paragraphs), keyword_index + section_range + 1)
        original_styles = _save_paragraph_styles(doc, start_index, len(new_content))
        _replace_paragraphs_in_range(doc, start_index, end_index, new_content, original_styles)
        return f"Replaced paragraphs containing keyword '{keyword}' and their surrounding content, keeping original format and style"
    except Exception as e:
        return f"Failed to replace content: {str(e)}"


def _handle_add_list_item(params: dict) -> str:
    text = params.get("text")
    if text is None:
        return "Missing required parameter: text"
    list_type = params.get("list_type", "bullet")
    level = params.get("level", 0)
    if not isinstance(level, int) or not (0 <= level <= 8):
        return "Invalid level: must be an integer between 0 and 8"
    try:
        if not processor.current_document:
            return "No document is open"
        style = 'List Bullet' if list_type == "bullet" else 'List Number'
        paragraph = processor.current_document.add_paragraph(text, style=style)
        if level > 0:
            paragraph.paragraph_format.left_indent = Cm(1.27 * level)
        return f"Added {list_type} list item"
    except Exception as e:
        return f"Failed to add list item: {str(e)}"


def _handle_add_image(params: dict) -> str:
    image_path = params.get("image_path")
    if not image_path:
        return "Missing required parameter: image_path"
    width = params.get("width")
    height = params.get("height")
    try:
        if not processor.current_document:
            return "No document is open"
        if not os.path.exists(image_path):
            return f"Image file does not exist: {image_path}"
        kwargs = {}
        if width is not None:
            kwargs['width'] = Cm(width)
        if height is not None:
            kwargs['height'] = Cm(height)
        processor.current_document.add_picture(image_path, **kwargs)
        return f"Image inserted: {image_path}"
    except Exception as e:
        return f"Failed to insert image: {str(e)}"


def _handle_set_header(params: dict) -> str:
    text = params.get("text")
    if text is None:
        return "Missing required parameter: text"
    section_index = params.get("section_index", 0)
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        if section_index < 0 or section_index >= len(doc.sections):
            return f"Section index out of range: {section_index}, document has {len(doc.sections)} sections"
        section = doc.sections[section_index]
        header = section.header
        header.is_linked_to_previous = False
        header.paragraphs[0].text = text
        return f"Header set for section {section_index}"
    except Exception as e:
        return f"Failed to set header: {str(e)}"


def _handle_set_footer(params: dict) -> str:
    text = params.get("text")
    if text is None:
        return "Missing required parameter: text"
    section_index = params.get("section_index", 0)
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        if section_index < 0 or section_index >= len(doc.sections):
            return f"Section index out of range: {section_index}, document has {len(doc.sections)} sections"
        section = doc.sections[section_index]
        footer = section.footer
        footer.is_linked_to_previous = False
        footer.paragraphs[0].text = text
        return f"Footer set for section {section_index}"
    except Exception as e:
        return f"Failed to set footer: {str(e)}"


def _handle_set_page_orientation(params: dict) -> str:
    orientation = params.get("orientation")
    if orientation is None:
        return "Missing required parameter: orientation"
    section_index = params.get("section_index", 0)
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        if section_index < 0 or section_index >= len(doc.sections):
            return f"Section index out of range: {section_index}, document has {len(doc.sections)} sections"
        section = doc.sections[section_index]
        if orientation == "portrait":
            section.orientation = WD_ORIENT.PORTRAIT
        elif orientation == "landscape":
            section.orientation = WD_ORIENT.LANDSCAPE
        else:
            return f"Invalid orientation: {orientation}. Use 'portrait' or 'landscape'"
        return f"Page orientation set to {orientation} for section {section_index}"
    except Exception as e:
        return f"Failed to set page orientation: {str(e)}"


def _handle_set_page_size(params: dict) -> str:
    width = params.get("width")
    height = params.get("height")
    preset = params.get("preset")
    section_index = params.get("section_index", 0)
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        if section_index < 0 or section_index >= len(doc.sections):
            return f"Section index out of range: {section_index}, document has {len(doc.sections)} sections"
        section = doc.sections[section_index]
        presets = {
            "A4": (21.0, 29.7),
            "A3": (29.7, 42.0),
            "Letter": (21.59, 27.94),
            "Legal": (21.59, 35.56),
        }
        if preset:
            if preset not in presets:
                return f"Unknown preset: {preset}. Available: {', '.join(presets.keys())}"
            w, h = presets[preset]
        elif width is not None and height is not None:
            w, h = width, height
        else:
            return "Please specify either preset or both width and height"
        section.page_width = Cm(w)
        section.page_height = Cm(h)
        return f"Page size set to {w}x{h} cm for section {section_index}"
    except Exception as e:
        return f"Failed to set page size: {str(e)}"


def _handle_add_table_column(params: dict) -> str:
    table_index = params.get("table_index")
    col_index = params.get("col_index")
    data = params.get("data")
    if table_index is None or col_index is None:
        return "Missing required parameters: table_index, col_index"
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        if not doc.tables:
            return "No tables in document"
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Table index out of range: {table_index}, document has {len(doc.tables)} tables"
        table = doc.tables[table_index]
        num_rows = len(table.rows)
        if col_index == -1:
            col_index = len(table.columns)
        if col_index < 0 or col_index > len(table.columns):
            return f"Column index out of range: {col_index}"
        tbl = table._tbl
        for i, tr in enumerate(tbl.findall(qn('w:tr'))):
            tc = OxmlElement('w:tc')
            tc_p = OxmlElement('w:p')
            tc.append(tc_p)
            existing_tcs = tr.findall(qn('w:tc'))
            if col_index < len(existing_tcs):
                existing_tcs[col_index].addprevious(tc)
            else:
                tr.append(tc)
            if data and i < len(data):
                tc_p.text = str(data[i])
        return f"Added column at index {col_index} to table {table_index}"
    except Exception as e:
        return f"Failed to add table column: {str(e)}"


def _handle_delete_table_column(params: dict) -> str:
    table_index = params.get("table_index")
    col_index = params.get("col_index")
    if table_index is None or col_index is None:
        return "Missing required parameters: table_index, col_index"
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        if not doc.tables:
            return "No tables in document"
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Table index out of range: {table_index}, document has {len(doc.tables)} tables"
        table = doc.tables[table_index]
        if col_index < 0 or col_index >= len(table.columns):
            return f"Column index out of range: {col_index}, table has {len(table.columns)} columns"
        tbl = table._tbl
        for tr in tbl.findall(qn('w:tr')):
            tcs = tr.findall(qn('w:tc'))
            if col_index < len(tcs):
                tr.remove(tcs[col_index])
        return f"Deleted column {col_index} from table {table_index}"
    except Exception as e:
        return f"Failed to delete table column: {str(e)}"


def _handle_set_line_spacing(params: dict) -> str:
    spacing = params.get("spacing")
    paragraph_index = params.get("paragraph_index")
    if spacing is None:
        return "Missing required parameter: spacing"
    valid_prefixes = ("multiple:", "exact:")
    if spacing not in ("1.0", "1.5", "2.0") and not spacing.startswith(valid_prefixes):
        return f"Invalid spacing: '{spacing}'. Use '1.0', '1.5', '2.0', 'multiple:N', or 'exact:N'"
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        if spacing.startswith("multiple:"):
            value = float(spacing.split(":")[1])
            rule = WD_LINE_SPACING.MULTIPLE
        elif spacing.startswith("exact:"):
            value = Pt(float(spacing.split(":")[1]))
            rule = WD_LINE_SPACING.EXACTLY
        else:
            value = float(spacing)
            rule = WD_LINE_SPACING.MULTIPLE
        if paragraph_index is not None:
            if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
                return f"Paragraph index out of range: {paragraph_index}"
            paragraphs = [doc.paragraphs[paragraph_index]]
        else:
            paragraphs = doc.paragraphs
        for para in paragraphs:
            para.paragraph_format.line_spacing = value
            para.paragraph_format.line_spacing_rule = rule
        target = f"paragraph {paragraph_index}" if paragraph_index is not None else "all paragraphs"
        return f"Line spacing set to {spacing} for {target}"
    except Exception as e:
        return f"Failed to set line spacing: {str(e)}"


def _handle_add_hyperlink(params: dict) -> str:
    paragraph_index = params.get("paragraph_index")
    text = params.get("text")
    url = params.get("url")
    if paragraph_index is None or text is None or url is None:
        return "Missing required parameters: paragraph_index, text, url"
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Paragraph index out of range: {paragraph_index}, document has {len(doc.paragraphs)} paragraphs"
        paragraph = doc.paragraphs[paragraph_index]
        part = paragraph.part
        r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        c = OxmlElement('w:color')
        c.set(qn('w:val'), '0563C1')
        rPr.append(c)
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
        new_run.append(rPr)
        t = OxmlElement('w:t')
        t.text = text
        new_run.append(t)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
        return f"Hyperlink added to paragraph {paragraph_index}"
    except Exception as e:
        return f"Failed to add hyperlink: {str(e)}"


def _handle_set_text_highlight(params: dict) -> str:
    paragraph_index = params.get("paragraph_index")
    color = params.get("color")
    if paragraph_index is None or color is None:
        return "Missing required parameters: paragraph_index, color"
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Paragraph index out of range: {paragraph_index}"
        paragraph = doc.paragraphs[paragraph_index]
        color_upper = color.upper()
        if not hasattr(WD_COLOR_INDEX, color_upper):
            return f"Invalid color: {color}. Available: YELLOW, GREEN, CYAN, MAGENTA, BLUE, RED, DARK_BLUE, DARK_CYAN, DARK_GREEN, DARK_MAGENTA, DARK_RED, DARK_YELLOW, LIGHT_GRAY, DARK_GRAY, BLACK, WHITE"
        highlight = WD_COLOR_INDEX[color_upper]
        for run in paragraph.runs:
            run.font.highlight_color = highlight
        return f"Highlight set to {color} for paragraph {paragraph_index}"
    except Exception as e:
        return f"Failed to set highlight: {str(e)}"


def _handle_set_text_strikethrough(params: dict) -> str:
    paragraph_index = params.get("paragraph_index")
    double = params.get("double", False)
    if paragraph_index is None:
        return "Missing required parameter: paragraph_index"
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Paragraph index out of range: {paragraph_index}"
        paragraph = doc.paragraphs[paragraph_index]
        for run in paragraph.runs:
            if double:
                run.font.double_strike = True
            else:
                run.font.strikethrough = True
        style = "double strikethrough" if double else "strikethrough"
        return f"{style} set for paragraph {paragraph_index}"
    except Exception as e:
        return f"Failed to set strikethrough: {str(e)}"


def _handle_add_section(params: dict) -> str:
    orientation = params.get("orientation")
    page_width = params.get("page_width")
    page_height = params.get("page_height")
    top_margin = params.get("top_margin")
    bottom_margin = params.get("bottom_margin")
    left_margin = params.get("left_margin")
    right_margin = params.get("right_margin")
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        section = doc.add_section(WD_SECTION_START.NEW_PAGE)
        if orientation == "portrait":
            section.orientation = WD_ORIENT.PORTRAIT
        elif orientation == "landscape":
            section.orientation = WD_ORIENT.LANDSCAPE
        if page_width is not None:
            section.page_width = Cm(page_width)
        if page_height is not None:
            section.page_height = Cm(page_height)
        if top_margin is not None:
            section.top_margin = Cm(top_margin)
        if bottom_margin is not None:
            section.bottom_margin = Cm(bottom_margin)
        if left_margin is not None:
            section.left_margin = Cm(left_margin)
        if right_margin is not None:
            section.right_margin = Cm(right_margin)
        return f"New section added (section {len(doc.sections) - 1})"
    except Exception as e:
        return f"Failed to add section: {str(e)}"


def _handle_add_bookmark(params: dict) -> str:
    paragraph_index = params.get("paragraph_index")
    bookmark_name = params.get("bookmark_name")
    if paragraph_index is None or bookmark_name is None:
        return "Missing required parameters: paragraph_index, bookmark_name"
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Paragraph index out of range: {paragraph_index}"
        paragraph = doc.paragraphs[paragraph_index]
        bookmark_id = 0
        for p in doc.paragraphs:
            for bm in p._element.findall(qn('w:bookmarkStart')):
                try:
                    bid = int(bm.get(qn('w:id'), 0))
                    if bid >= bookmark_id:
                        bookmark_id = bid + 1
                except ValueError:
                    pass
        bookmark_start = OxmlElement('w:bookmarkStart')
        bookmark_start.set(qn('w:id'), str(bookmark_id))
        bookmark_start.set(qn('w:name'), bookmark_name)
        bookmark_end = OxmlElement('w:bookmarkEnd')
        bookmark_end.set(qn('w:id'), str(bookmark_id))
        paragraph._p.insert(0, bookmark_start)
        paragraph._p.append(bookmark_end)
        return f"Bookmark '{bookmark_name}' added to paragraph {paragraph_index}"
    except Exception as e:
        return f"Failed to add bookmark: {str(e)}"


def _handle_add_comment(params: dict) -> str:
    paragraph_index = params.get("paragraph_index")
    text = params.get("text")
    author = params.get("author", "MCP Server")
    if paragraph_index is None or text is None:
        return "Missing required parameters: paragraph_index, text"
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Paragraph index out of range: {paragraph_index}"
        paragraph = doc.paragraphs[paragraph_index]
        from docx.opc.part import Part
        from docx.opc.packuri import PackURI
        import docx.opc.constants as opc_constants
        comments_part_name = '/word/comments.xml'
        part = doc.part
        if not hasattr(part, '_comments_part') or part._comments_part is None:
            comments_xml = (
                '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                '<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
                ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
                '</w:comments>'
            )
            comments_part = Part(
                PackURI(comments_part_name),
                'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml',
                comments_xml.encode('utf-8'),
                doc.part.package
            )
            doc.part.relate_to(comments_part, opc_constants.RELATIONSHIP_TYPE.COMMENTS)
            part._comments_part = comments_part
        comments_part = part._comments_part
        from lxml import etree
        comments_elem = etree.fromstring(comments_part.blob)
        comment_id = 0
        for c in comments_elem.findall(qn('w:comment')):
            try:
                cid = int(c.get(qn('w:id'), 0))
                if cid >= comment_id:
                    comment_id = cid + 1
            except ValueError:
                pass
        comment = OxmlElement('w:comment')
        comment.set(qn('w:id'), str(comment_id))
        comment.set(qn('w:author'), author)
        comment.set(qn('w:date'), '2025-01-01T00:00:00Z')
        comment_p = OxmlElement('w:p')
        comment_r = OxmlElement('w:r')
        comment_t = OxmlElement('w:t')
        comment_t.text = text
        comment_r.append(comment_t)
        comment_p.append(comment_r)
        comment.append(comment_p)
        comments_elem.append(comment)
        comments_part._blob = etree.tostring(comments_elem, xml_declaration=True, encoding='UTF-8', standalone=True)
        comment_range_start = OxmlElement('w:commentRangeStart')
        comment_range_start.set(qn('w:id'), str(comment_id))
        comment_range_end = OxmlElement('w:commentRangeEnd')
        comment_range_end.set(qn('w:id'), str(comment_id))
        comment_reference_run = OxmlElement('w:r')
        comment_reference_rPr = OxmlElement('w:rPr')
        comment_reference = OxmlElement('w:commentReference')
        comment_reference.set(qn('w:id'), str(comment_id))
        comment_reference_rPr.append(comment_reference)
        comment_reference_run.append(comment_reference_rPr)
        paragraph._p.insert(0, comment_range_start)
        paragraph._p.append(comment_range_end)
        paragraph._p.append(comment_reference_run)
        return f"Comment added to paragraph {paragraph_index}"
    except Exception as e:
        return f"Failed to add comment: {str(e)}"


def _handle_add_footnote(params: dict) -> str:
    paragraph_index = params.get("paragraph_index")
    text = params.get("text")
    if paragraph_index is None or text is None:
        return "Missing required parameters: paragraph_index, text"
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Paragraph index out of range: {paragraph_index}"
        paragraph = doc.paragraphs[paragraph_index]
        from docx.opc.part import Part
        from docx.opc.packuri import PackURI
        import docx.opc.constants as opc_constants
        part = doc.part
        if not hasattr(part, '_footnotes_part') or part._footnotes_part is None:
            footnotes_xml = (
                '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
                ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
                '<w:footnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:footnote>'
                '<w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>'
                '</w:footnotes>'
            )
            footnotes_part = Part(
                PackURI('/word/footnotes.xml'),
                'application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml',
                footnotes_xml.encode('utf-8'),
                doc.part.package
            )
            doc.part.relate_to(footnotes_part, opc_constants.RELATIONSHIP_TYPE.FOOTNOTES)
            part._footnotes_part = footnotes_part
        footnotes_part = part._footnotes_part
        from lxml import etree
        footnotes_elem = etree.fromstring(footnotes_part.blob)
        footnote_id = 1
        for fn in footnotes_elem.findall(qn('w:footnote')):
            try:
                fid = int(fn.get(qn('w:id'), 0))
                if fid >= footnote_id:
                    footnote_id = fid + 1
            except ValueError:
                pass
        footnote = OxmlElement('w:footnote')
        footnote.set(qn('w:id'), str(footnote_id))
        fn_p = OxmlElement('w:p')
        fn_r = OxmlElement('w:r')
        fn_t = OxmlElement('w:t')
        fn_t.text = text
        fn_r.append(fn_t)
        fn_p.append(fn_r)
        footnote.append(fn_p)
        footnotes_elem.append(footnote)
        footnotes_part._blob = etree.tostring(footnotes_elem, xml_declaration=True, encoding='UTF-8', standalone=True)
        run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rStyle = OxmlElement('w:rStyle')
        rStyle.set(qn('w:val'), 'FootnoteReference')
        rPr.append(rStyle)
        run.append(rPr)
        footnote_ref = OxmlElement('w:footnoteReference')
        footnote_ref.set(qn('w:id'), str(footnote_id))
        run.append(footnote_ref)
        paragraph._p.append(run)
        return f"Footnote {footnote_id} added to paragraph {paragraph_index}"
    except Exception as e:
        return f"Failed to add footnote: {str(e)}"


def _handle_add_endnote(params: dict) -> str:
    paragraph_index = params.get("paragraph_index")
    text = params.get("text")
    if paragraph_index is None or text is None:
        return "Missing required parameters: paragraph_index, text"
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Paragraph index out of range: {paragraph_index}"
        paragraph = doc.paragraphs[paragraph_index]
        from docx.opc.part import Part
        from docx.opc.packuri import PackURI
        import docx.opc.constants as opc_constants
        part = doc.part
        if not hasattr(part, '_endnotes_part') or part._endnotes_part is None:
            endnotes_xml = (
                '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                '<w:endnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
                ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
                '<w:endnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:endnote>'
                '<w:endnote w:type="continuationSeparator" w:id="0"><w:p><w:r><w:continuationSeparator/></w:r></w:p></w:endnote>'
                '</w:endnotes>'
            )
            endnotes_part = Part(
                PackURI('/word/endnotes.xml'),
                'application/vnd.openxmlformats-officedocument.wordprocessingml.endnotes+xml',
                endnotes_xml.encode('utf-8'),
                doc.part.package
            )
            doc.part.relate_to(endnotes_part, opc_constants.RELATIONSHIP_TYPE.ENDNOTES)
            part._endnotes_part = endnotes_part
        endnotes_part = part._endnotes_part
        from lxml import etree
        endnotes_elem = etree.fromstring(endnotes_part.blob)
        endnote_id = 1
        for en in endnotes_elem.findall(qn('w:endnote')):
            try:
                eid = int(en.get(qn('w:id'), 0))
                if eid >= endnote_id:
                    endnote_id = eid + 1
            except ValueError:
                pass
        endnote = OxmlElement('w:endnote')
        endnote.set(qn('w:id'), str(endnote_id))
        en_p = OxmlElement('w:p')
        en_r = OxmlElement('w:r')
        en_t = OxmlElement('w:t')
        en_t.text = text
        en_r.append(en_t)
        en_p.append(en_r)
        endnote.append(en_p)
        endnotes_elem.append(endnote)
        endnotes_part._blob = etree.tostring(endnotes_elem, xml_declaration=True, encoding='UTF-8', standalone=True)
        run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rStyle = OxmlElement('w:rStyle')
        rStyle.set(qn('w:val'), 'EndnoteReference')
        rPr.append(rStyle)
        run.append(rPr)
        endnote_ref = OxmlElement('w:endnoteReference')
        endnote_ref.set(qn('w:id'), str(endnote_id))
        run.append(endnote_ref)
        paragraph._p.append(run)
        return f"Endnote {endnote_id} added to paragraph {paragraph_index}"
    except Exception as e:
        return f"Failed to add endnote: {str(e)}"


def _handle_add_table_of_contents(params: dict) -> str:
    title = params.get("title", "目录")
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        if title:
            doc.add_paragraph(title, style='Heading 1')
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar_begin)
        run2 = paragraph.add_run()
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
        run2._r.append(instrText)
        run3 = paragraph.add_run()
        fldChar_separate = OxmlElement('w:fldChar')
        fldChar_separate.set(qn('w:fldCharType'), 'separate')
        run3._r.append(fldChar_separate)
        run4 = paragraph.add_run('(请在 Word 中右键点击此处，选择"更新域"以生成目录)')
        run4.font.color.rgb = RGBColor(128, 128, 128)
        run5 = paragraph.add_run()
        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')
        run5._r.append(fldChar_end)
        return "Table of contents inserted (update in Word to populate)"
    except Exception as e:
        return f"Failed to insert table of contents: {str(e)}"


def _handle_set_superscript(params: dict) -> str:
    paragraph_index = params.get("paragraph_index")
    if paragraph_index is None:
        return "Missing required parameter: paragraph_index"
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Paragraph index out of range: {paragraph_index}"
        for run in doc.paragraphs[paragraph_index].runs:
            run.font.superscript = True
        return f"Superscript set for paragraph {paragraph_index}"
    except Exception as e:
        return f"Failed to set superscript: {str(e)}"


def _handle_set_subscript(params: dict) -> str:
    paragraph_index = params.get("paragraph_index")
    if paragraph_index is None:
        return "Missing required parameter: paragraph_index"
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Paragraph index out of range: {paragraph_index}"
        for run in doc.paragraphs[paragraph_index].runs:
            run.font.subscript = True
        return f"Subscript set for paragraph {paragraph_index}"
    except Exception as e:
        return f"Failed to set subscript: {str(e)}"


def _handle_add_tab_stop(params: dict) -> str:
    paragraph_index = params.get("paragraph_index")
    position = params.get("position")
    alignment = params.get("alignment", "left")
    if paragraph_index is None or position is None:
        return "Missing required parameters: paragraph_index, position"
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Paragraph index out of range: {paragraph_index}"
        from docx.enum.text import WD_TAB_ALIGNMENT
        align_map = {
            "left": WD_TAB_ALIGNMENT.LEFT,
            "center": WD_TAB_ALIGNMENT.CENTER,
            "right": WD_TAB_ALIGNMENT.RIGHT,
            "decimal": WD_TAB_ALIGNMENT.DECIMAL,
        }
        if alignment not in align_map:
            return f"Invalid alignment: {alignment}. Use left, center, right, or decimal"
        paragraph = doc.paragraphs[paragraph_index]
        paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(position), align_map[alignment])
        return f"Tab stop added at {position} cm ({alignment}) to paragraph {paragraph_index}"
    except Exception as e:
        return f"Failed to add tab stop: {str(e)}"


def _handle_set_text_direction(params: dict) -> str:
    paragraph_index = params.get("paragraph_index")
    direction = params.get("direction")
    if paragraph_index is None or direction is None:
        return "Missing required parameters: paragraph_index, direction"
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Paragraph index out of range: {paragraph_index}"
        paragraph = doc.paragraphs[paragraph_index]
        pPr = paragraph._element.get_or_add_pPr()
        if direction == "rtl":
            bidi = OxmlElement('w:bidi')
            bidi.set(qn('w:val'), '1')
            pPr.append(bidi)
        elif direction == "ltr":
            bidi = pPr.find(qn('w:bidi'))
            if bidi is not None:
                pPr.remove(bidi)
        else:
            return f"Invalid direction: {direction}. Use 'ltr' or 'rtl'"
        return f"Text direction set to {direction} for paragraph {paragraph_index}"
    except Exception as e:
        return f"Failed to set text direction: {str(e)}"


def _handle_get_paragraph(params: dict) -> str:
    paragraph_index = params.get("paragraph_index")
    if paragraph_index is None:
        return "Missing required parameter: paragraph_index"
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Paragraph index out of range: {paragraph_index}, document has {len(doc.paragraphs)} paragraphs"
        paragraph = doc.paragraphs[paragraph_index]
        alignment_map = {
            WD_PARAGRAPH_ALIGNMENT.LEFT: "left",
            WD_PARAGRAPH_ALIGNMENT.CENTER: "center",
            WD_PARAGRAPH_ALIGNMENT.RIGHT: "right",
            WD_PARAGRAPH_ALIGNMENT.JUSTIFY: "justify",
        }
        alignment = alignment_map.get(paragraph.alignment, str(paragraph.alignment)) if paragraph.alignment else None
        runs_info = []
        for run in paragraph.runs:
            run_info = {
                "text": run.text,
                "bold": run.bold,
                "italic": run.italic,
                "underline": run.underline,
                "font_size": str(run.font.size) if run.font.size else None,
                "font_name": run.font.name,
                "color": str(run.font.color.rgb) if run.font.color and run.font.color.rgb else None,
            }
            runs_info.append(run_info)
        result = {
            "text": paragraph.text,
            "style": paragraph.style.name if paragraph.style else None,
            "alignment": alignment,
            "runs": runs_info,
        }
        return json.dumps(result, ensure_ascii=False, indent=2)
    except Exception as e:
        return f"Failed to get paragraph: {str(e)}"


def _handle_get_table_cell(params: dict) -> str:
    table_index = params.get("table_index")
    row_index = params.get("row_index")
    col_index = params.get("col_index")
    if table_index is None or row_index is None or col_index is None:
        return "Missing required parameters: table_index, row_index, col_index"
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        if not doc.tables:
            return "No tables in document"
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Table index out of range: {table_index}, document has {len(doc.tables)} tables"
        table = doc.tables[table_index]
        if row_index < 0 or row_index >= len(table.rows):
            return f"Row index out of range: {row_index}, table has {len(table.rows)} rows"
        if col_index < 0 or col_index >= len(table.columns):
            return f"Column index out of range: {col_index}, table has {len(table.columns)} columns"
        cell = table.cell(row_index, col_index)
        return json.dumps({"text": cell.text}, ensure_ascii=False)
    except Exception as e:
        return f"Failed to get table cell: {str(e)}"


def _handle_get_page_info(params: dict) -> str:
    section_index = params.get("section_index", 0)
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        if section_index < 0 or section_index >= len(doc.sections):
            return f"Section index out of range: {section_index}, document has {len(doc.sections)} sections"
        section = doc.sections[section_index]
        orient_map = {WD_ORIENT.PORTRAIT: "portrait", WD_ORIENT.LANDSCAPE: "landscape"}
        result = {
            "orientation": orient_map.get(section.orientation, str(section.orientation)),
            "page_width": round(section.page_width.cm, 2) if section.page_width else None,
            "page_height": round(section.page_height.cm, 2) if section.page_height else None,
            "top_margin": round(section.top_margin.cm, 2) if section.top_margin else None,
            "bottom_margin": round(section.bottom_margin.cm, 2) if section.bottom_margin else None,
            "left_margin": round(section.left_margin.cm, 2) if section.left_margin else None,
            "right_margin": round(section.right_margin.cm, 2) if section.right_margin else None,
        }
        return json.dumps(result, ensure_ascii=False, indent=2)
    except Exception as e:
        return f"Failed to get page info: {str(e)}"


def _handle_set_paragraph_text(params: dict) -> str:
    paragraph_index = params.get("paragraph_index")
    text = params.get("text")
    if paragraph_index is None or text is None:
        return "Missing required parameters: paragraph_index, text"
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Paragraph index out of range: {paragraph_index}, document has {len(doc.paragraphs)} paragraphs"
        paragraph = doc.paragraphs[paragraph_index]
        # Save format from first run
        fmt = {}
        if paragraph.runs:
            run0 = paragraph.runs[0]
            fmt["bold"] = run0.bold
            fmt["italic"] = run0.italic
            fmt["underline"] = run0.underline
            fmt["font_size"] = run0.font.size
            fmt["font_name"] = run0.font.name
            fmt["color"] = run0.font.color.rgb if run0.font.color and run0.font.color.rgb else None
        # Clear and rewrite
        paragraph.clear()
        run = paragraph.add_run(text)
        if fmt:
            run.bold = fmt["bold"]
            run.italic = fmt["italic"]
            run.underline = fmt["underline"]
            if fmt["font_size"]:
                run.font.size = fmt["font_size"]
            if fmt["font_name"]:
                _set_run_font(run, fmt["font_name"])
            if fmt["color"]:
                run.font.color.rgb = fmt["color"]
        return f"Paragraph {paragraph_index} text updated"
    except Exception as e:
        return f"Failed to set paragraph text: {str(e)}"


def _handle_set_paragraph_format(params: dict) -> str:
    paragraph_index = params.get("paragraph_index")
    if paragraph_index is None:
        return "Missing required parameter: paragraph_index"
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Paragraph index out of range: {paragraph_index}, document has {len(doc.paragraphs)} paragraphs"
        paragraph = doc.paragraphs[paragraph_index]
        bold = params.get("bold")
        italic = params.get("italic")
        underline = params.get("underline")
        font_size = params.get("font_size")
        font_name = params.get("font_name")
        color = params.get("color")
        alignment = params.get("alignment")
        if color is not None and (not color.startswith('#') or len(color) != 7):
            return f"Invalid color format: '{color}'. Use '#RRGGBB' format (e.g. '#FF0000')"
        for run in paragraph.runs:
            if bold is not None:
                run.bold = bold
            if italic is not None:
                run.italic = italic
            if underline is not None:
                run.underline = underline
            if font_size is not None:
                run.font.size = Pt(font_size)
            if font_name is not None:
                _set_run_font(run, font_name)
            if color is not None and color.startswith('#') and len(color) == 7:
                run.font.color.rgb = _parse_hex_color(color)
        if alignment:
            align_map = {
                "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
                "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
                "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
                "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
            }
            if alignment in align_map:
                paragraph.alignment = align_map[alignment]
        return f"Format updated for paragraph {paragraph_index}"
    except Exception as e:
        return f"Failed to set paragraph format: {str(e)}"


def _handle_set_cell_format(params: dict) -> str:
    table_index = params.get("table_index")
    row_index = params.get("row_index")
    col_index = params.get("col_index")
    if table_index is None or row_index is None or col_index is None:
        return "Missing required parameters: table_index, row_index, col_index"
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        if not doc.tables:
            return "No tables in document"
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Table index out of range: {table_index}, document has {len(doc.tables)} tables"
        table = doc.tables[table_index]
        if row_index < 0 or row_index >= len(table.rows):
            return f"Row index out of range: {row_index}, table has {len(table.rows)} rows"
        if col_index < 0 or col_index >= len(table.columns):
            return f"Column index out of range: {col_index}, table has {len(table.columns)} columns"
        cell = table.cell(row_index, col_index)
        bold = params.get("bold")
        italic = params.get("italic")
        underline = params.get("underline")
        font_size = params.get("font_size")
        font_name = params.get("font_name")
        color = params.get("color")
        if color is not None and (not color.startswith('#') or len(color) != 7):
            return f"Invalid color format: '{color}'. Use '#RRGGBB' format (e.g. '#FF0000')"
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                if bold is not None:
                    run.bold = bold
                if italic is not None:
                    run.italic = italic
                if underline is not None:
                    run.underline = underline
                if font_size is not None:
                    run.font.size = Pt(font_size)
                if font_name is not None:
                    _set_run_font(run, font_name)
                if color is not None and color.startswith('#') and len(color) == 7:
                    run.font.color.rgb = _parse_hex_color(color)
        return f"Format updated for cell ({row_index},{col_index}) in table {table_index}"
    except Exception as e:
        return f"Failed to set cell format: {str(e)}"


def _handle_set_paragraph_spacing(params: dict) -> str:
    paragraph_index = params.get("paragraph_index")
    space_before = params.get("space_before")
    space_after = params.get("space_after")
    first_line_indent = params.get("first_line_indent")
    if paragraph_index is None:
        return "Missing required parameter: paragraph_index"
    if space_before is None and space_after is None and first_line_indent is None:
        return "At least one of space_before, space_after, first_line_indent is required"
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Paragraph index out of range: {paragraph_index}, document has {len(doc.paragraphs)} paragraphs"
        paragraph = doc.paragraphs[paragraph_index]
        if space_before is not None:
            paragraph.paragraph_format.space_before = Pt(space_before)
        if space_after is not None:
            paragraph.paragraph_format.space_after = Pt(space_after)
        if first_line_indent is not None:
            paragraph.paragraph_format.first_line_indent = Cm(first_line_indent)
        return f"Spacing updated for paragraph {paragraph_index}"
    except Exception as e:
        return f"Failed to set paragraph spacing: {str(e)}"


def _handle_add_page_number(params: dict) -> str:
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        section_index = params.get("section_index", 0)
        if section_index < 0 or section_index >= len(doc.sections):
            return f"Section index out of range: {section_index}, document has {len(doc.sections)} sections"
        section = doc.sections[section_index]
        footer = section.footer
        footer.is_linked_to_previous = False
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.clear()
        run1 = paragraph.add_run()
        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')
        run1._r.append(fldChar_begin)
        run2 = paragraph.add_run()
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' PAGE '
        run2._r.append(instrText)
        run3 = paragraph.add_run()
        fldChar_separate = OxmlElement('w:fldChar')
        fldChar_separate.set(qn('w:fldCharType'), 'separate')
        run3._r.append(fldChar_separate)
        run4 = paragraph.add_run('1')
        run5 = paragraph.add_run()
        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')
        run5._r.append(fldChar_end)
        return f"Page number field added to footer of section {section_index}"
    except Exception as e:
        return f"Failed to add page number: {str(e)}"


def _handle_set_table_borders(params: dict) -> str:
    table_index = params.get("table_index")
    color = params.get("color", "000000")
    size = params.get("size", 4)
    if table_index is None:
        return "Missing required parameter: table_index"
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        if not doc.tables:
            return "No tables in document"
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Table index out of range: {table_index}, document has {len(doc.tables)} tables"
        table = doc.tables[table_index]
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
        borders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), str(size))
            border.set(qn('w:color'), color)
            border.set(qn('w:space'), '0')
            borders.append(border)
        existing = tblPr.find(qn('w:tblBorders'))
        if existing is not None:
            tblPr.remove(existing)
        tblPr.append(borders)
        return f"Borders set for table {table_index}"
    except Exception as e:
        return f"Failed to set table borders: {str(e)}"


def _handle_set_cell_shading(params: dict) -> str:
    table_index = params.get("table_index")
    row_index = params.get("row_index")
    col_index = params.get("col_index")
    color = params.get("color")
    if table_index is None or row_index is None or col_index is None or color is None:
        return "Missing required parameters: table_index, row_index, col_index, color"
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        if not doc.tables:
            return "No tables in document"
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Table index out of range: {table_index}"
        table = doc.tables[table_index]
        if row_index < 0 or row_index >= len(table.rows):
            return f"Row index out of range: {row_index}"
        if col_index < 0 or col_index >= len(table.columns):
            return f"Column index out of range: {col_index}"
        cell = table.cell(row_index, col_index)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color)
        existing = tcPr.find(qn('w:shd'))
        if existing is not None:
            tcPr.remove(existing)
        tcPr.append(shd)
        return f"Shading set for cell ({row_index},{col_index}) in table {table_index}"
    except Exception as e:
        return f"Failed to set cell shading: {str(e)}"


def _handle_set_paragraph_border(params: dict) -> str:
    paragraph_index = params.get("paragraph_index")
    if paragraph_index is None:
        return "Missing required parameter: paragraph_index"
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Paragraph index out of range: {paragraph_index}"
        paragraph = doc.paragraphs[paragraph_index]
        pPr = paragraph._element.get_or_add_pPr()
        borders = OxmlElement('w:pBdr')
        for side in ['top', 'left', 'bottom', 'right']:
            val = params.get(f"{side}_style")
            if val:
                border = OxmlElement(f'w:{side}')
                border.set(qn('w:val'), val)
                border.set(qn('w:sz'), str(params.get(f"{side}_size", 4)))
                border.set(qn('w:color'), params.get(f"{side}_color", "000000"))
                border.set(qn('w:space'), str(params.get(f"{side}_space", 1)))
                borders.append(border)
        if len(borders) == 0:
            return "At least one border side is required (top_style, left_style, bottom_style, right_style)"
        existing = pPr.find(qn('w:pBdr'))
        if existing is not None:
            pPr.remove(existing)
        pPr.append(borders)
        return f"Border set for paragraph {paragraph_index}"
    except Exception as e:
        return f"Failed to set paragraph border: {str(e)}"


def _handle_set_paragraph_shading(params: dict) -> str:
    paragraph_index = params.get("paragraph_index")
    color = params.get("color")
    if paragraph_index is None or color is None:
        return "Missing required parameters: paragraph_index, color"
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Paragraph index out of range: {paragraph_index}"
        paragraph = doc.paragraphs[paragraph_index]
        pPr = paragraph._element.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color)
        existing = pPr.find(qn('w:shd'))
        if existing is not None:
            pPr.remove(existing)
        pPr.append(shd)
        return f"Shading set for paragraph {paragraph_index}"
    except Exception as e:
        return f"Failed to set paragraph shading: {str(e)}"


def _handle_set_columns(params: dict) -> str:
    num_columns = params.get("num_columns")
    if num_columns is None:
        return "Missing required parameter: num_columns"
    section_index = params.get("section_index", 0)
    space = params.get("space")
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        if section_index < 0 or section_index >= len(doc.sections):
            return f"Section index out of range: {section_index}"
        section = doc.sections[section_index]
        sectPr = section._sectPr
        cols = OxmlElement('w:cols')
        cols.set(qn('w:num'), str(num_columns))
        if space is not None:
            cols.set(qn('w:space'), str(Cm(space)))
        existing = sectPr.find(qn('w:cols'))
        if existing is not None:
            sectPr.remove(existing)
        sectPr.append(cols)
        return f"Columns set to {num_columns} for section {section_index}"
    except Exception as e:
        return f"Failed to set columns: {str(e)}"


def _handle_create_style(params: dict) -> str:
    style_name = params.get("style_name")
    style_type = params.get("style_type", "paragraph")
    if style_name is None:
        return "Missing required parameter: style_name"
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        type_map = {
            "paragraph": WD_STYLE_TYPE.PARAGRAPH,
            "character": WD_STYLE_TYPE.CHARACTER,
            "table": WD_STYLE_TYPE.TABLE,
            "list": WD_STYLE_TYPE.LIST,
        }
        if style_type not in type_map:
            return f"Invalid style_type: {style_type}. Use: {', '.join(type_map.keys())}"
        for existing in doc.styles:
            if existing.name == style_name:
                return f"Style '{style_name}' already exists"
        style = doc.styles.add_style(style_name, type_map[style_type])
        bold = params.get("bold")
        italic = params.get("italic")
        font_size = params.get("font_size")
        font_name = params.get("font_name")
        color = params.get("color")
        if bold is not None:
            style.font.bold = bold
        if italic is not None:
            style.font.italic = italic
        if font_size is not None:
            style.font.size = Pt(font_size)
        if font_name is not None:
            style.font.name = font_name
        if color and color.startswith('#') and len(color) == 7:
            style.font.color.rgb = _parse_hex_color(color)
        return f"Style '{style_name}' created ({style_type})"
    except Exception as e:
        return f"Failed to create style: {str(e)}"


def _handle_modify_style(params: dict) -> str:
    style_name = params.get("style_name")
    if style_name is None:
        return "Missing required parameter: style_name"
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        style = None
        for s in doc.styles:
            if s.name == style_name:
                style = s
                break
        if style is None:
            return f"Style not found: '{style_name}'"
        bold = params.get("bold")
        italic = params.get("italic")
        underline = params.get("underline")
        font_size = params.get("font_size")
        font_name = params.get("font_name")
        color = params.get("color")
        if bold is not None:
            style.font.bold = bold
        if italic is not None:
            style.font.italic = italic
        if underline is not None:
            style.font.underline = underline
        if font_size is not None:
            style.font.size = Pt(font_size)
        if font_name is not None:
            style.font.name = font_name
        if color and color.startswith('#') and len(color) == 7:
            style.font.color.rgb = _parse_hex_color(color)
        return f"Style '{style_name}' modified"
    except Exception as e:
        return f"Failed to modify style: {str(e)}"


def _handle_list_styles(params: dict) -> str:
    try:
        if not processor.current_document:
            return "No document is open"
        doc = processor.current_document
        style_type = params.get("style_type")
        type_map = {
            "paragraph": WD_STYLE_TYPE.PARAGRAPH,
            "character": WD_STYLE_TYPE.CHARACTER,
            "table": WD_STYLE_TYPE.TABLE,
            "list": WD_STYLE_TYPE.LIST,
        }
        styles = []
        for style in doc.styles:
            if style_type and style.type != type_map.get(style_type):
                continue
            styles.append(f"{style.name} ({style.type.name if style.type else 'unknown'})")
        if not styles:
            return "No styles found"
        return f"Found {len(styles)} styles:\n" + "\n".join(styles[:30])
    except Exception as e:
        return f"Failed to list styles: {str(e)}"


# ============================================================================
# Route registry
# ============================================================================

ROUTE_HANDLERS = {
    # Document management
    "create_document": _handle_create_document,
    "open_document": _handle_open_document,
    "save_document": _handle_save_document,
    "save_as_document": _handle_save_as_document,
    "create_document_copy": _handle_create_document_copy,
    "close_document": _handle_close_document,
    "reload_document": _handle_reload_document,
    # Content addition
    "add_paragraph": _handle_add_paragraph,
    "add_heading": _handle_add_heading,
    "add_table": _handle_add_table,
    "add_list_item": _handle_add_list_item,
    "add_image": _handle_add_image,
    "add_page_break": _handle_add_page_break,
    "add_section": _handle_add_section,
    "add_table_of_contents": _handle_add_table_of_contents,
    # Content editing
    "search_text": _handle_search_text,
    "search_and_replace": _handle_search_and_replace,
    "find_and_replace": _handle_find_and_replace,
    "replace_section": _handle_replace_section,
    "edit_section_by_keyword": _handle_edit_section_by_keyword,
    "delete_paragraph": _handle_delete_paragraph,
    "delete_text": _handle_delete_text,
    "edit_table_cell": _handle_edit_table_cell,
    # Table operations
    "add_table_row": _handle_add_table_row,
    "add_table_column": _handle_add_table_column,
    "delete_table_row": _handle_delete_table_row,
    "delete_table_column": _handle_delete_table_column,
    "merge_table_cells": _handle_merge_table_cells,
    "split_table": _handle_split_table,
    # Formatting
    "set_page_margins": _handle_set_page_margins,
    "set_page_orientation": _handle_set_page_orientation,
    "set_page_size": _handle_set_page_size,
    "set_header": _handle_set_header,
    "set_footer": _handle_set_footer,
    "set_line_spacing": _handle_set_line_spacing,
    "set_text_highlight": _handle_set_text_highlight,
    "set_text_strikethrough": _handle_set_text_strikethrough,
    "set_superscript": _handle_set_superscript,
    "set_subscript": _handle_set_subscript,
    "set_text_direction": _handle_set_text_direction,
    "add_tab_stop": _handle_add_tab_stop,
    "add_hyperlink": _handle_add_hyperlink,
    # Annotations and references
    "add_bookmark": _handle_add_bookmark,
    "add_comment": _handle_add_comment,
    "add_footnote": _handle_add_footnote,
    "add_endnote": _handle_add_endnote,
    # Query
    "get_document_info": _handle_get_document_info,
    "get_paragraph": _handle_get_paragraph,
    "get_table_cell": _handle_get_table_cell,
    "get_page_info": _handle_get_page_info,
    # Paragraph and cell modification
    "set_paragraph_text": _handle_set_paragraph_text,
    "set_paragraph_format": _handle_set_paragraph_format,
    "set_cell_format": _handle_set_cell_format,
    # P1: Paragraph spacing and indent
    "set_paragraph_spacing": _handle_set_paragraph_spacing,
    # P1: Page number
    "add_page_number": _handle_add_page_number,
    # P1: Table formatting
    "set_table_borders": _handle_set_table_borders,
    "set_cell_shading": _handle_set_cell_shading,
    # P1: Paragraph borders and shading
    "set_paragraph_border": _handle_set_paragraph_border,
    "set_paragraph_shading": _handle_set_paragraph_shading,
    # P1: Column layout
    "set_columns": _handle_set_columns,
    # P1: Style management
    "create_style": _handle_create_style,
    "modify_style": _handle_modify_style,
    "list_styles": _handle_list_styles,
}


# ============================================================================
# Gateway tool
# ============================================================================

@mcp.tool()
def docx_process(ctx: Context, route: str, params: Dict[str, Any]) -> str:
    """
    IMPORTANT: Before calling this tool, you MUST first invoke the "docx-process" skill to load the complete parameter specifications. Do NOT call this tool without first loading the skill documentation.

    Word document processing gateway. Use route to select the operation.

    Parameters:
    - route: Operation name (e.g. "create_document", "add_paragraph", "search_text")
    - params: Parameters for the selected operation, see skill documentation for details
    """
    handler = ROUTE_HANDLERS.get(route)
    if handler is None:
        available = ", ".join(sorted(ROUTE_HANDLERS.keys()))
        return f"Unknown route: '{route}'. Available routes: {available}"
    return handler(params)


if __name__ == "__main__":
    mcp.run()
